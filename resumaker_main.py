"""TODO: какой нахуй туду? нам лишь бы проект в школке защитить и забить на этот говнокод

если нет уже существующей базы данных программа начинает писать файл с новой строки

использовать docxtpl вместо python-docx чтобы делать резюме красивее
								и иметь возможность менять шаблоны резюме на ходу

fix "extensions must match" error

свой title bar через frameless window
"""

import qdarkstyle
import ctypes
from sys import argv
from os import getenv
import fix_qt_import_error
from PyQt5 import QtWidgets
import resumaker_ui
from docx import Document
from docx.shared import Inches


myappid = u'mycompany.myproduct.subproduct.version'

ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

photopath = ("", "")


class Resumaker(QtWidgets.QMainWindow, resumaker_ui.Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.savebutton.clicked.connect(self.browse_folder)
        self.Addphoto.clicked.connect(self.addphoto)

    def addphoto(self):
        global photopath
        photopath = QtWidgets.QFileDialog.getOpenFileName(self, "Приложить личное фото", getenv('USERPROFILE'),
                                                          "Файлы изображений (*.png *.jpg )")
        print(photopath)

    def browse_folder(self):
        directory = QtWidgets.QFileDialog.getSaveFileName(self, "Сохранение резюме",
                                                          getenv('USERPROFILE') + r"\Desktop",
                                                          'Документ Word (*.docx)')
        if directory[0] != "":
            resume_savefile = Document()
            resume_savefile.save(directory[0])

            resume = Document(docx=directory[0])

            sex = self.sex_comboBox.currentIndex()
            sex_text = self.sex_comboBox.currentText()
            second_name = self.secondname.displayText()
            name = self.name.displayText()
            patronymic = self.patronymic.displayText()
            post = self.post.displayText()
            phone_number = self.phoneNumber.displayText()
            professional_skills = self.professional_skills.toPlainText()
            email = self.email.displayText()
            birth_date = self.birthDate.date().toPyDate()
            personal_qualities = self.personal_qualities.toPlainText()
            edu_level = self.edu_level.currentText()
            experience = self.experience.toPlainText()
            birth_year = int(str(birth_date)[0:4])
            birth_month = str(birth_date)[5:7]
            birth_day = str(birth_date)[8:10]

            birth_date = birth_day + "." + birth_month + "." + str(birth_year)
            age = str(2019 - birth_year)

            if sex == 1:
                sex = "Мужчина, "
            else:
                sex = "Женщина, "

            resume.add_heading('Резюме', level=0)
            resume.add_heading(text=str(second_name) + " " + str(name) + " " + str(patronymic), level=1)
            resume.add_heading(sex + age + " лет, дата рождения: " + birth_date, level=1)

            if photopath[0] != "":
                resume.add_picture(photopath[0], width=Inches(1.25))

            p = resume.add_paragraph(text="", style=None)
            p.add_run('Должность: ').bold = True
            p.add_run(post)

            if phone_number != "":
                mobile_phone = resume.add_paragraph(text="", style=None)
                mobile_phone.add_run('Моб. телефон: ').bold = True
                mobile_phone.add_run(phone_number)

            if email != "":
                em = resume.add_paragraph(text="", style=None)
                em.add_run('E-mail: ').bold = True
                em.add_run(email)

            edu = resume.add_paragraph(text="", style=None)
            edu.add_run('Уровень образования: ').bold = True
            edu.add_run(edu_level)

            if personal_qualities != "":
                personal_qualities_title = resume.add_paragraph(text="", style=None)
                personal_qualities_title.add_run('Персональные качества: ').bold = True

                resume.add_paragraph(text=personal_qualities, style=None)

            if professional_skills != "":
                professional_skills_title = resume.add_paragraph(text="", style=None)
                professional_skills_title.add_run('Профессиональные навыки: ').bold = True

                resume.add_paragraph(text=professional_skills, style=None)

            if experience != "":
                experience_title = resume.add_paragraph(text="", style=None)
                experience_title.add_run('Опыт работы:').bold = True

                resume.add_paragraph(text=experience, style=None)

            # database_file = open('resumaker_database.rdb', 'a', encoding="utf-8")
            # write_string = "\n" + name + "\t" + second_name + "\t" + post + "\t" + sex_text + "\t" + age +\
            #                "\t" + edu_level
            # database_file.write(write_string)
            # database_file.close()

            db = open('resumaker_database.rdb', 'r+', encoding="utf-8")
            if db.read() != "":
            	db.write("\n")
            db.write(name + "," + second_name + "," + post + "," +\
            		 age + "," + edu_level + "," + directory[0])
            db.close()

            resume.save(directory[0])


def main():
    app = QtWidgets.QApplication(argv)
    window = Resumaker()
    app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt5()) # темная тема
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
